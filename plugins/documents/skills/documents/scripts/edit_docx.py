#!/usr/bin/env python3
"""Apply conservative text, metadata, and append operations to a DOCX copy."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any, Iterable, Mapping

from docx import Document

from document_builder import (
    apply_core_properties,
    load_json_object,
    object_list,
    render_blocks,
)
from docx_layout import preset_by_name


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, type=Path, help="Source .docx file.")
    parser.add_argument("--operations", required=True, type=Path, help="UTF-8 JSON edit operations.")
    parser.add_argument("--output", required=True, type=Path, help="Destination .docx copy.")
    parser.add_argument("--report", type=Path, help="Optional JSON report path.")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    input_path = args.input.expanduser().resolve()
    operations_path = args.operations.expanduser().resolve()
    output_path = args.output.expanduser().resolve()
    if input_path.suffix.lower() != ".docx" or output_path.suffix.lower() != ".docx":
        raise ValueError("--input and --output must be .docx files.")
    if input_path == output_path:
        raise ValueError("The editor preserves its source; choose a different --output path.")
    if not input_path.is_file():
        raise FileNotFoundError(f"Input document not found: {input_path}")
    if not operations_path.is_file():
        raise FileNotFoundError(f"Operations file not found: {operations_path}")

    operations = load_json_object(operations_path)
    document = Document(input_path)
    replacement_report = apply_replacements(document, operations.get("replacements"))

    properties = operations.get("properties")
    if properties is not None:
        if not isinstance(properties, Mapping):
            raise ValueError("properties must be an object.")
        title = str(properties.get("title") or document.core_properties.title or input_path.stem)
        apply_core_properties(document, title, properties, preserve_missing_author=True)

    append_blocks = operations.get("appendBlocks")
    appended = 0
    if append_blocks is not None:
        blocks = object_list(append_blocks, "appendBlocks")
        preset = preset_by_name(operations.get("preset"))
        render_blocks(document, blocks, preset, base_dir=operations_path.parent)
        appended = len(blocks)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    report = {
        "input": str(input_path),
        "output": str(output_path),
        "replacements": replacement_report,
        "appendedBlocks": appended,
        "size": output_path.stat().st_size,
    }
    if args.report:
        report_path = args.report.expanduser().resolve()
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False))
    return 0


def apply_replacements(document: Any, raw_replacements: Any) -> list[dict[str, Any]]:
    if raw_replacements is None:
        return []
    replacements = object_list(raw_replacements, "replacements")
    paragraphs = list(iter_document_paragraphs(document))
    report: list[dict[str, Any]] = []
    for index, operation in enumerate(replacements):
        find = required_string(operation.get("find"), f"replacements[{index}].find")
        replace = str(operation.get("replace") or "")
        case_sensitive = operation.get("caseSensitive") is not False
        requested_count = non_negative_integer(operation.get("count"), 0)
        required = operation.get("required") is not False
        changed = 0
        for paragraph in paragraphs:
            remaining = 0 if requested_count == 0 else requested_count - changed
            if requested_count and remaining <= 0:
                break
            changed += replace_in_paragraph(
                paragraph,
                find,
                replace,
                count=remaining,
                case_sensitive=case_sensitive,
            )
        if required and changed == 0:
            raise ValueError(f"Required replacement did not match: {find!r}.")
        report.append({"find": find, "changed": changed, "requested": requested_count or "all"})
    return report


def replace_in_paragraph(
    paragraph: Any,
    needle: str,
    replacement: str,
    *,
    count: int,
    case_sensitive: bool,
) -> int:
    runs = list(paragraph.runs)
    if not runs:
        return 0
    text = "".join(run.text for run in runs)
    positions: list[int] = []
    if case_sensitive:
        cursor = 0
        while cursor <= len(text) - len(needle):
            position = text.find(needle, cursor)
            if position < 0:
                break
            positions.append(position)
            if count and len(positions) >= count:
                break
            cursor = position + len(needle)
    else:
        for match in re.finditer(re.escape(needle), text, flags=re.IGNORECASE):
            positions.append(match.start())
            if count and len(positions) >= count:
                break
    if not positions:
        return 0

    for position in reversed(positions):
        replace_run_range(runs, position, position + len(needle), replacement)
    return len(positions)


def replace_run_range(runs: list[Any], start: int, end: int, replacement: str) -> None:
    start_run, start_offset = run_position(runs, start, allow_end=False)
    end_run, end_offset = run_position(runs, end, allow_end=True)
    if start_run == end_run:
        text = runs[start_run].text
        runs[start_run].text = text[:start_offset] + replacement + text[end_offset:]
        return

    start_text = runs[start_run].text
    end_text = runs[end_run].text
    runs[start_run].text = start_text[:start_offset] + replacement
    for index in range(start_run + 1, end_run):
        runs[index].text = ""
    runs[end_run].text = end_text[end_offset:]


def run_position(runs: list[Any], position: int, *, allow_end: bool) -> tuple[int, int]:
    cursor = 0
    for index, run in enumerate(runs):
        next_cursor = cursor + len(run.text)
        if position < next_cursor or (allow_end and position == next_cursor):
            return index, position - cursor
        cursor = next_cursor
    return len(runs) - 1, len(runs[-1].text)


def iter_document_paragraphs(document: Any) -> Iterable[Any]:
    yield from iter_container_paragraphs(document)
    seen_parts: set[int] = set()
    for section in document.sections:
        for container in (section.header, section.footer):
            part_id = id(container.part)
            if part_id in seen_parts:
                continue
            seen_parts.add(part_id)
            yield from iter_container_paragraphs(container)


def iter_container_paragraphs(container: Any) -> Iterable[Any]:
    yield from container.paragraphs
    seen_cells: set[int] = set()
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                yield from iter_container_paragraphs(cell)


def required_string(value: Any, label: str) -> str:
    if not isinstance(value, str) or not value:
        raise ValueError(f"{label} must be a non-empty string.")
    return value


def non_negative_integer(value: Any, default: int) -> int:
    if value is None:
        return default
    try:
        result = int(value)
    except (TypeError, ValueError) as error:
        raise ValueError("Replacement count must be a non-negative integer.") from error
    if result < 0:
        raise ValueError("Replacement count must be a non-negative integer.")
    return result


if __name__ == "__main__":
    raise SystemExit(main())
