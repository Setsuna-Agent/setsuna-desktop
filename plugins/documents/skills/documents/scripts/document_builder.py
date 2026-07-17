"""Render the plugin's JSON content model into a python-docx document."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Iterable, Mapping, Sequence

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from docx_layout import (
    DocumentPreset,
    apply_run_font,
    configure_new_document,
    ensure_support_styles,
    normalize_hex_color,
    paragraph_alignment,
    preset_by_name,
    set_cell_fill,
    set_cell_margins,
    set_paragraph_panel,
    set_paragraph_rule,
    set_repeat_table_header,
    set_table_geometry,
    usable_width_dxa,
)


CALLOUT_TONES = {
    "info": ("EAF2F7", "315E7D"),
    "success": ("EAF5EF", "2F7A55"),
    "warning": ("FFF4DF", "B46B12"),
    "neutral": ("F1F3F5", "66717E"),
}


def load_json_object(path: Path) -> dict[str, Any]:
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as error:
        raise ValueError(f"Invalid JSON in {path}: {error}") from error
    if not isinstance(value, dict):
        raise ValueError(f"Expected one JSON object in {path}.")
    return value


def build_document(specification: Mapping[str, Any], base_dir: Path) -> Any:
    title = required_text(specification.get("title"), "title")
    preset = preset_by_name(specification.get("preset"))
    document = Document()
    configure_new_document(document, preset, optional_text(specification.get("pageSize")))
    apply_core_properties(document, title, specification.get("properties"))
    add_front_matter(document, title, specification, preset)
    render_blocks(
        document,
        object_list(specification.get("blocks"), "blocks"),
        preset,
        base_dir=base_dir,
    )
    return document


def apply_core_properties(
    document: Any,
    title: str,
    raw_properties: Any,
    *,
    preserve_missing_author: bool = False,
) -> None:
    properties = raw_properties if isinstance(raw_properties, Mapping) else {}
    core = document.core_properties
    core.title = title
    if optional_text(properties.get("subject")):
        core.subject = optional_text(properties.get("subject"))
    if optional_text(properties.get("author")):
        core.author = optional_text(properties.get("author"))
        core.last_modified_by = optional_text(properties.get("author"))
    elif not preserve_missing_author:
        core.author = ""
        core.last_modified_by = ""
    keywords = properties.get("keywords")
    if isinstance(keywords, list):
        core.keywords = ", ".join(str(value).strip() for value in keywords if str(value).strip())


def add_front_matter(
    document: Any,
    title: str,
    specification: Mapping[str, Any],
    preset: DocumentPreset,
) -> None:
    title_paragraph = document.add_paragraph(style="Setsuna Title")
    add_inline_content(title_paragraph, title, preset)

    subtitle = optional_text(specification.get("subtitle"))
    if subtitle:
        subtitle_paragraph = document.add_paragraph(style="Setsuna Subtitle")
        add_inline_content(subtitle_paragraph, subtitle, preset)

    metadata = specification.get("metadata")
    if metadata is not None:
        for index, item in enumerate(object_list(metadata, "metadata")):
            label = required_text(item.get("label"), f"metadata[{index}].label")
            value = required_text(item.get("value"), f"metadata[{index}].value")
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(1.5)
            label_run = paragraph.add_run(f"{label}: ")
            label_run.bold = True
            apply_run_font(label_run, preset, color=preset.muted, size_pt=9)
            value_run = paragraph.add_run(value)
            apply_run_font(value_run, preset, color=preset.muted, size_pt=9)
        document.add_paragraph().paragraph_format.space_after = Pt(2)


def render_blocks(
    document: Any,
    blocks: Sequence[Mapping[str, Any]],
    preset: DocumentPreset,
    *,
    base_dir: Path,
) -> None:
    ensure_support_styles(document, preset)
    for index, block in enumerate(blocks):
        block_type = str(block.get("type") or "").strip()
        if block_type == "heading":
            add_heading(document, block, preset, index)
        elif block_type == "paragraph":
            add_paragraph(document, block, preset)
        elif block_type in {"bullets", "numbered"}:
            add_list(document, block, preset, numbered=block_type == "numbered", index=index)
        elif block_type == "checklist":
            add_checklist(document, block, preset, index)
        elif block_type == "callout":
            add_callout(document, block, preset)
        elif block_type == "quote":
            add_quote(document, block, preset)
        elif block_type == "table":
            add_table(document, block, preset, index)
        elif block_type == "image":
            add_image(document, block, preset, base_dir, index)
        elif block_type == "rule":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            set_paragraph_rule(paragraph, preset.muted)
        elif block_type == "spacer":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(bounded_float(block.get("points"), 8, 0, 72))
        elif block_type == "pageBreak":
            document.add_page_break()
        else:
            raise ValueError(f"blocks[{index}].type is unsupported: {block_type!r}.")


def add_heading(document: Any, block: Mapping[str, Any], preset: DocumentPreset, index: int) -> None:
    level = integer(block.get("level"), 1)
    if level not in (1, 2, 3):
        raise ValueError(f"blocks[{index}].level must be 1, 2, or 3.")
    paragraph = document.add_paragraph(style=f"Heading {level}")
    add_inline_content(paragraph, block.get("text"), preset)


def add_paragraph(document: Any, block: Mapping[str, Any], preset: DocumentPreset) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = paragraph_alignment(block.get("align"), "left")
    add_inline_content(paragraph, block.get("text"), preset)


def add_list(
    document: Any,
    block: Mapping[str, Any],
    preset: DocumentPreset,
    *,
    numbered: bool,
    index: int,
) -> None:
    items = value_list(block.get("items"), f"blocks[{index}].items")
    for item in items:
        level = integer(item.get("level"), 0) if isinstance(item, Mapping) else 0
        if level < 0 or level > 2:
            raise ValueError("List item level must be between 0 and 2.")
        base_style = "List Number" if numbered else "List Bullet"
        style_name = base_style if level == 0 else f"{base_style} {level + 1}"
        paragraph = document.add_paragraph(style=style_name)
        paragraph.paragraph_format.space_after = Pt(max(1.5, preset.paragraph_after_pt / 2))
        add_inline_content(paragraph, item, preset)


def add_checklist(document: Any, block: Mapping[str, Any], preset: DocumentPreset, index: int) -> None:
    items = value_list(block.get("items"), f"blocks[{index}].items")
    for item in items:
        checked = bool(item.get("checked")) if isinstance(item, Mapping) else False
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.first_line_indent = Inches(-0.18)
        marker = paragraph.add_run("☒ " if checked else "☐ ")
        apply_run_font(marker, preset, color=preset.accent)
        add_inline_content(paragraph, item, preset)


def add_callout(document: Any, block: Mapping[str, Any], preset: DocumentPreset) -> None:
    tone = str(block.get("tone") or "info").strip().lower()
    fill, border = CALLOUT_TONES.get(tone, CALLOUT_TONES["neutral"])
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(9)
    set_paragraph_panel(paragraph, fill, border)
    label = optional_text(block.get("label"))
    if label:
        label_run = paragraph.add_run(f"{label}  ")
        label_run.bold = True
        apply_run_font(label_run, preset, color=border)
    add_inline_content(paragraph, block.get("text"), preset)


def add_quote(document: Any, block: Mapping[str, Any], preset: DocumentPreset) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.28)
    paragraph.paragraph_format.right_indent = Inches(0.2)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    set_paragraph_panel(paragraph, "F7F8FA", preset.muted)
    add_inline_content(paragraph, block.get("text"), preset, italic=True)
    attribution = optional_text(block.get("attribution"))
    if attribution:
        run = paragraph.add_run(f"  — {attribution}")
        run.italic = True
        apply_run_font(run, preset, color=preset.muted, size_pt=max(8, preset.body_size_pt - 1))


def add_table(document: Any, block: Mapping[str, Any], preset: DocumentPreset, index: int) -> None:
    columns = [required_text(value, f"blocks[{index}].columns") for value in value_list(block.get("columns"), f"blocks[{index}].columns")]
    if not columns:
        raise ValueError(f"blocks[{index}].columns cannot be empty.")
    raw_rows = value_list(block.get("rows"), f"blocks[{index}].rows")
    rows = [normalized_row(row, columns, index) for row in raw_rows]
    alignments = normalized_alignments(block.get("alignments"), len(columns))
    weights = normalized_width_weights(block.get("widths"), columns, rows)
    width_dxa = distribute_width(usable_width_dxa(document.sections[-1]), weights)

    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    header = table.rows[0]
    set_repeat_table_header(header)
    for column_index, (cell, label) in enumerate(zip(header.cells, columns)):
        populate_cell(cell, label, preset, bold=True, color="FFFFFF", align="center")
        set_cell_fill(cell, preset.table_header_fill)
        set_cell_margins(cell, preset.table_cell_margin_dxa)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for raw_row in rows:
        row = table.add_row()
        for column_index, (cell, value) in enumerate(zip(row.cells, raw_row)):
            populate_cell(cell, value, preset, align=alignments[column_index])
            set_cell_margins(cell, preset.table_cell_margin_dxa)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    set_table_geometry(table, width_dxa)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_image(
    document: Any,
    block: Mapping[str, Any],
    preset: DocumentPreset,
    base_dir: Path,
    index: int,
) -> None:
    raw_path = required_text(block.get("path"), f"blocks[{index}].path")
    image_path = Path(raw_path).expanduser()
    if not image_path.is_absolute():
        image_path = (base_dir / image_path).resolve()
    if not image_path.is_file():
        raise ValueError(f"Image does not exist: {image_path}")
    width_inches = bounded_float(block.get("widthInches"), 5.8, 0.5, 7.5)
    document.add_picture(str(image_path), width=Inches(width_inches))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = optional_text(block.get("caption"))
    if caption:
        paragraph = document.add_paragraph(style="Caption")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline_content(paragraph, caption, preset)


def add_inline_content(
    paragraph: Any,
    value: Any,
    preset: DocumentPreset,
    *,
    italic: bool = False,
) -> None:
    if isinstance(value, Mapping) and isinstance(value.get("runs"), list):
        for run_spec in value["runs"]:
            if not isinstance(run_spec, Mapping):
                raise ValueError("Inline runs must be objects.")
            add_styled_run(paragraph, run_spec, preset, inherited_italic=italic)
        return
    text = inline_text(value)
    run = paragraph.add_run(text)
    run.italic = italic
    apply_run_font(run, preset)


def add_styled_run(
    paragraph: Any,
    specification: Mapping[str, Any],
    preset: DocumentPreset,
    *,
    inherited_italic: bool,
) -> None:
    text = str(specification.get("text") or "")
    run = paragraph.add_run(text)
    run.bold = bool(specification.get("bold"))
    run.italic = inherited_italic or bool(specification.get("italic"))
    run.underline = bool(specification.get("underline"))
    color = optional_text(specification.get("color"))
    size = bounded_float(specification.get("size"), preset.body_size_pt, 6, 72)
    code = bool(specification.get("code"))
    apply_run_font(run, preset, color=normalize_hex_color(color) if color else None, size_pt=size, code=code)


def populate_cell(
    cell: Any,
    value: Any,
    preset: DocumentPreset,
    *,
    bold: bool = False,
    color: str | None = None,
    align: str = "left",
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = paragraph_alignment(align)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = max(1.0, preset.body_line_spacing - 0.05)
    if isinstance(value, Mapping) and isinstance(value.get("runs"), list):
        add_inline_content(paragraph, value, preset)
        if bold or color:
            for run in paragraph.runs:
                if bold:
                    run.bold = True
                if color:
                    run.font.color.rgb = RGBColor.from_string(normalize_hex_color(color))
        return
    run = paragraph.add_run(inline_text(value))
    run.bold = bold
    apply_run_font(run, preset, color=color)


def normalized_row(value: Any, columns: Sequence[str], block_index: int) -> list[Any]:
    if isinstance(value, Mapping):
        return [value.get(column, "") for column in columns]
    if isinstance(value, list):
        if len(value) != len(columns):
            raise ValueError(f"blocks[{block_index}] table row has {len(value)} cells; expected {len(columns)}.")
        return value
    raise ValueError(f"blocks[{block_index}] table rows must be arrays or objects.")


def normalized_alignments(value: Any, count: int) -> list[str]:
    if value is None:
        return ["left"] * count
    values = value_list(value, "table.alignments")
    if len(values) != count:
        raise ValueError("Table alignment count must match its column count.")
    result = [str(item).strip().lower() for item in values]
    for alignment in result:
        paragraph_alignment(alignment)
    return result


def normalized_width_weights(value: Any, columns: Sequence[str], rows: Sequence[Sequence[Any]]) -> list[float]:
    if value is not None:
        weights = [float(item) for item in value_list(value, "table.widths")]
        if len(weights) != len(columns) or any(weight <= 0 for weight in weights):
            raise ValueError("Table widths must be positive values matching the column count.")
        return weights
    weights: list[float] = []
    for index, column in enumerate(columns):
        lengths = [len(str(column))]
        lengths.extend(len(inline_text(row[index])) for row in rows)
        weights.append(float(min(8, max(1.4, max(lengths, default=1) / 8))))
    return weights


def distribute_width(total: int, weights: Sequence[float]) -> list[int]:
    weight_sum = sum(weights)
    widths = [max(1, int(total * weight / weight_sum)) for weight in weights]
    widths[-1] += total - sum(widths)
    return widths


def inline_text(value: Any) -> str:
    if isinstance(value, Mapping):
        if "text" in value:
            return str(value.get("text") or "")
        if isinstance(value.get("runs"), list):
            return "".join(str(run.get("text") or "") for run in value["runs"] if isinstance(run, Mapping))
    if value is None:
        return ""
    return str(value)


def object_list(value: Any, label: str) -> list[Mapping[str, Any]]:
    values = value_list(value, label)
    if not all(isinstance(item, Mapping) for item in values):
        raise ValueError(f"{label} must contain only objects.")
    return list(values)


def value_list(value: Any, label: str) -> list[Any]:
    if not isinstance(value, list):
        raise ValueError(f"{label} must be an array.")
    return value


def required_text(value: Any, label: str) -> str:
    text = optional_text(value)
    if not text:
        raise ValueError(f"{label} must be a non-empty string.")
    return text


def optional_text(value: Any) -> str | None:
    if value is None:
        return None
    text = str(value).strip()
    return text or None


def integer(value: Any, default: int) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def bounded_float(value: Any, default: float, minimum: float, maximum: float) -> float:
    try:
        number = float(value)
    except (TypeError, ValueError):
        number = default
    return min(maximum, max(minimum, number))
